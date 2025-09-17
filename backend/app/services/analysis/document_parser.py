import PyPDF2
from docx import Document
import io
from typing import Dict, Optional
from loguru import logger
import httpx


class DocumentParser:
    """Parser for PDF and Word documents"""
    
    async def parse_document_from_url(self, url: str) -> Dict:
        """Download and parse document from URL"""
        try:
            # Download the document
            async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
                response = await client.get(url)
                response.raise_for_status()
                
                content_type = response.headers.get('content-type', '').lower()
                
                if 'pdf' in content_type:
                    return self._parse_pdf(response.content)
                elif 'wordprocessingml' in content_type or url.endswith('.docx'):
                    return self._parse_docx(response.content)
                else:
                    return {
                        "error": f"Unsupported document type: {content_type}",
                        "document_type": "unknown"
                    }
                    
        except Exception as e:
            logger.error(f"Error parsing document from {url}: {e}")
            return {
                "error": str(e),
                "document_type": "unknown"
            }
    
    def _parse_pdf(self, content: bytes) -> Dict:
        """Extract text from PDF"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num}: {e}")
                    continue
            
            # Clean up the text
            text = text.strip()
            text = ' '.join(text.split())  # Normalize whitespace
            
            return {
                "content": text,
                "word_count": len(text.split()),
                "page_count": len(pdf_reader.pages),
                "document_type": "pdf",
                "error": None
            }
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            return {
                "error": str(e),
                "document_type": "pdf",
                "content": "",
                "word_count": 0
            }
    
    def _parse_docx(self, content: bytes) -> Dict:
        """Extract text from Word document"""
        try:
            doc = Document(io.BytesIO(content))
            
            # Extract paragraph text
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # Extract table text
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        tables_text.append("\t".join(row_text))
            
            # Combine all text
            all_text = "\n\n".join(paragraphs)
            if tables_text:
                all_text += "\n\n" + "\n".join(tables_text)
            
            return {
                "content": all_text.strip(),
                "word_count": len(all_text.split()),
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
                "document_type": "docx",
                "error": None
            }
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            return {
                "error": str(e),
                "document_type": "docx",
                "content": "",
                "word_count": 0
            }
    
    def parse_local_file(self, file_path: str) -> Dict:
        """Parse a local document file"""
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
            
            if file_path.lower().endswith('.pdf'):
                return self._parse_pdf(content)
            elif file_path.lower().endswith('.docx'):
                return self._parse_docx(content)
            else:
                return {
                    "error": "Unsupported file type",
                    "document_type": "unknown"
                }
        except Exception as e:
            logger.error(f"Error parsing local file {file_path}: {e}")
            return {
                "error": str(e),
                "document_type": "unknown"
            } 